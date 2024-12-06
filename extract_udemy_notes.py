from bs4 import BeautifulSoup
import pandas as pd
from fpdf import FPDF
import re
import xlsxwriter
import html2text
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def clean_html_with_bold_formatting(html_content):
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Replace <b> and <strong> tags with markers for bold text
    for bold in soup.find_all(['b', 'strong']):
        bold.insert_before('**')  # Markdown-style bold marker
        bold.insert_after('**')
    
    for line_break in soup.find_all('br'):
        line_break.replace_with('\n')  # Preserve line breaks
    
    return soup.get_text().strip()

def extract_notes_from_html(html_file):
    with open(html_file, 'r', encoding='utf-8') as file:
        soup = BeautifulSoup(file, 'html.parser')
    
    notes_containers = soup.find_all('div', class_='lecture-bookmark-v2--content-container--hoogx')
    notes_data = []
    for note in notes_containers:
        # Extract content and preserve bold formatting
        content_container = note.find('div', class_='rt-scaffolding')
        if content_container:
            content_html = str(content_container)
            content = clean_html_with_bold_formatting(content_html)  # Process bold text
        else:
            content = 'No Content'
        
        # Extract primary title
        primary_title = note.find_previous('div', class_='lecture-bookmark-v2--section--j0ti8')
        primary_title_text = primary_title.text.strip() if primary_title else 'No Primary Title'
        
        # Extract secondary title
        secondary_title = note.find_previous('div', class_='ud-text-sm')
        secondary_title_text = secondary_title.text.strip() if secondary_title else 'No Secondary Title'
        
        # Extract timestamp
        timestamp_container = note.find_previous('span', id=lambda x: x and x.startswith('bookmark-'))
        timestamp = timestamp_container.text.strip() if timestamp_container else 'No Time'
        
        notes_data.append({
            'Timestamp': timestamp,
            'Primary Title': primary_title_text,
            'Secondary Title': secondary_title_text,
            'Content': content
        })
    
    return notes_data


def save_to_excel(notes_data, excel_file):
    import xlsxwriter

    # Create a new workbook and add a worksheet
    workbook = xlsxwriter.Workbook(excel_file)
    worksheet = workbook.add_worksheet()

    # Write headers
    headers = ['Timestamp', 'Primary Title', 'Secondary Title', 'Content']
    for col_num, header in enumerate(headers):
        worksheet.write(0, col_num, header)

    # Formats for wrapped text and bold text
    wrap_format = workbook.add_format({'text_wrap': True})
    bold_format = workbook.add_format({'bold': True})

    # Write notes to worksheet
    row = 1
    for note in notes_data:
        worksheet.write(row, 0, note['Timestamp'], wrap_format)
        worksheet.write(row, 1, note['Primary Title'], wrap_format)
        worksheet.write(row, 2, note['Secondary Title'], wrap_format)

        # Process content for bold markers
        content_parts = note['Content'].split("**")  # Split content by '**'
        rich_text = []

        for i, part in enumerate(content_parts):
            if i % 2 == 0:  # Regular text
                if part.strip():  # Ignore empty fragments
                    rich_text.append(part)
            else:  # Bold text
                if part.strip():  # Ignore empty fragments
                    rich_text.append(bold_format)
                    rich_text.append(part)

        # Clean up invalid sequences
        valid_rich_text = []
        for i, fragment in enumerate(rich_text):
            # Skip consecutive formats or empty strings
            if isinstance(fragment, str) or (i > 0 and isinstance(fragment, xlsxwriter.format.Format)):
                valid_rich_text.append(fragment)

        # Validate final rich text
        if len(valid_rich_text) >= 3 and isinstance(valid_rich_text[0], str):  # At least one text and one format
            worksheet.write_rich_string(row, 3, *valid_rich_text, wrap_format)
        else:
            # Fallback: Write plain text
            plain_text = ''.join(content_parts)  # Join all parts as plain text
            worksheet.write(row, 3, plain_text, wrap_format)

        row += 1

    workbook.close()
    print(f"Data saved to Excel: {excel_file}")



def export_to_pdf(notes_data, pdf_file):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    def render_text_with_bold(pdf, text):
        parts = text.split("**")  # Split on bold markers
        bold = False
        for part in parts:
            if bold:
                pdf.set_font("Arial", style="B", size=12)  # Bold
            else:
                pdf.set_font("Arial", style="", size=12)  # Regular
            pdf.multi_cell(0, 10, txt=part)  # Handle each part as a new line
            bold = not bold  # Toggle bold mode

    for note in notes_data:
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(0, 10, txt=f"{note['Timestamp']}", ln=True)
        pdf.cell(0, 10, txt=note['Primary Title'], ln=True)
        pdf.cell(0, 10, txt=note['Secondary Title'], ln=True)
        
        # Render content with bold
        render_text_with_bold(pdf, note['Content'])
        pdf.ln(10)  # Add spacing between notes

    pdf.output(pdf_file)
    print(f"Data exported to PDF: {pdf_file}")

def save_to_word(notes_data, word_file):
    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading("Udemy Notes", level=1)

    # Iterate through notes and add them to the Word document
    for note in notes_data:
        # Add timestamp
        if note['Timestamp']:
            doc.add_paragraph(f"Timestamp: {note['Timestamp']}", style="Heading 2")

        # Add primary and secondary titles
        if note['Primary Title']:
            doc.add_paragraph(note['Primary Title'], style="Heading 3")
        if note['Secondary Title']:
            doc.add_paragraph(note['Secondary Title'], style="Heading 3")

        # Add content with bold formatting
        content_parts = note['Content'].split("**")  # Split content by '**'
        para = doc.add_paragraph()
        for i, part in enumerate(content_parts):
            if i % 2 == 0:  # Regular text
                para.add_run(part)
            else:  # Bold text
                para.add_run(part).bold = True

        # Add spacing between notes
        doc.add_paragraph("\n")

    # Save the document
    doc.save(word_file)
    print(f"Data saved to Word: {word_file}")


# Main function to automate the process
def main():
    html_file = r"C:\Users\tayfu\Desktop\Udemy Notes Extractor\course.html"

    excel_file = "Udemy_Notes.xlsx"
    pdf_file = "Udemy_Notes.pdf"
    word_file = "Udemy_Notes.docx"

    # Step 1: Extract notes from HTML
    notes_data = extract_notes_from_html(html_file)

    # Step 2: Reverse the order
    notes_data.reverse()

    # Step 3: Save to Excel
    save_to_excel(notes_data, excel_file)

    # Step 4: Export to PDF
    export_to_pdf(notes_data, pdf_file)

    # Step 5: Save to Word
    save_to_word(notes_data, word_file)


if __name__ == "__main__":
    main()

